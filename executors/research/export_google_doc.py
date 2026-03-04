#!/usr/bin/env python3
"""
Google Docs Exporter
====================
Exports a markdown file to a Google Doc with native formatting.

Strategy: creates a .docx first (proper formatting), then uploads to Google Drive
which auto-converts it to a Google Doc. This ensures perfect formatting.

Usage:
    python3 executors/research/export_google_doc.py <script_md_path> [--title "Title"] [--credentials path]
    python3 executors/research/export_google_doc.py --help

Arguments:
    script_md_path    Path to the markdown file to export

Options:
    --title TITLE           Document title (default: first # heading in the file)
    --credentials PATH      Path to OAuth credentials.json (default: ./credentials.json)

Installation:
    pip install python-docx google-api-python-client google-auth-httplib2 google-auth-oauthlib

Output JSON (stdout):
    {"success": true, "doc_url": "https://docs.google.com/document/d/.../edit", "doc_id": "...", "title": "..."}

Notes:
    - First run opens a browser for OAuth consent (token is cached afterward)
    - Token cached at ~/.cache/youtube-assistant/google_token.json
    - Scopes: drive.file (minimal — only accesses files this app creates)
    - Converts markdown to .docx first, then uploads to Google Drive as a Google Doc
"""
from __future__ import annotations

import json
import os
import re
import sys
import tempfile

SCOPES = [
    "https://www.googleapis.com/auth/drive.file",
]

TOKEN_DIR = os.path.expanduser("~/.cache/youtube-assistant")
TOKEN_PATH = os.path.join(TOKEN_DIR, "google_token.json")


def print_help() -> None:
    print(__doc__.strip())


def fail(message: str) -> None:
    """Print error JSON to stdout and exit 1."""
    print(json.dumps({"error": message}))
    sys.exit(1)


def check_dependencies() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        fail(
            "python-docx not found. Install with:\n"
            "  pip install python-docx"
        )
    try:
        import google.oauth2.credentials  # noqa: F401
        import google_auth_oauthlib.flow  # noqa: F401
        import googleapiclient.discovery  # noqa: F401
    except ImportError:
        fail(
            "Google API libraries not found. Install with:\n"
            "  pip install google-api-python-client google-auth-httplib2 google-auth-oauthlib"
        )


def authenticate(credentials_path: str):
    """Authenticate with Google OAuth2, returning credentials."""
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow

    creds = None

    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)

    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
        except Exception:
            creds = None

    if not creds or not creds.valid:
        if not os.path.exists(credentials_path):
            fail(
                f"credentials.json not found at {credentials_path}. To set up:\n"
                "  1. Go to https://console.cloud.google.com/\n"
                "  2. Create or select a project\n"
                "  3. Enable the Google Drive API\n"
                "  4. Create OAuth 2.0 credentials (Desktop application)\n"
                "  5. Download the JSON and save it as credentials.json in the project root\n"
                "  (See credentials.sample.json for the expected format)"
            )
        flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
        try:
            creds = flow.run_local_server(port=0)
        except Exception:
            try:
                creds = flow.run_console()
            except Exception as e:
                fail(f"OAuth authentication failed: {e}")

    os.makedirs(TOKEN_DIR, exist_ok=True)
    with open(TOKEN_PATH, "w") as f:
        f.write(creds.to_json())

    return creds


# ---------------------------------------------------------------------------
# Markdown parsing & Word document creation (same as export_word.py)
# ---------------------------------------------------------------------------

def parse_inline(text: str) -> list[dict]:
    """Parse inline markdown formatting into runs."""
    runs: list[dict] = []
    pattern = re.compile(
        r'\[([^\]]+)\]\(([^)]+)\)'
        r'|(\*{1,3})(.*?)\3'
    )
    pos = 0

    for match in pattern.finditer(text):
        if match.start() > pos:
            plain = text[pos:match.start()]
            if plain:
                runs.append({"text": plain, "bold": False, "italic": False, "url": None})

        if match.group(1) is not None:
            runs.append({
                "text": match.group(1),
                "bold": False, "italic": False,
                "url": match.group(2),
            })
        else:
            stars = len(match.group(3))
            inner = match.group(4)
            runs.append({
                "text": inner,
                "bold": stars >= 2,
                "italic": stars == 1 or stars == 3,
                "url": None,
            })
        pos = match.end()

    if pos < len(text):
        remaining = text[pos:]
        if remaining:
            runs.append({"text": remaining, "bold": False, "italic": False, "url": None})

    if not runs:
        runs.append({"text": text, "bold": False, "italic": False, "url": None})

    return runs


def parse_markdown(content: str) -> list[dict]:
    """Parse markdown into a list of blocks.

    Special handling:
    - [Source: ...] lines are appended to the preceding paragraph with a line break
    - First H1 heading and the metadata line (**Format**: ...) are skipped
    """
    blocks: list[dict] = []
    lines = content.split("\n")
    paragraph_lines: list[str] = []
    skipped_first_h1 = False

    def flush_paragraph():
        if paragraph_lines:
            combined_parts = []
            for j, pline in enumerate(paragraph_lines):
                if pline.endswith("\\"):
                    combined_parts.append(pline[:-1].rstrip())
                    combined_parts.append("\n")
                else:
                    combined_parts.append(pline)
                    if j < len(paragraph_lines) - 1:
                        combined_parts.append(" ")
            text = "".join(combined_parts)
            blocks.append({"type": "paragraph", "runs": parse_inline(text)})
            paragraph_lines.clear()

    for line in lines:
        stripped = line.strip()

        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            flush_paragraph()
            blocks.append({"type": "horizontal_rule"})
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            # Skip the first H1 (document title)
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                continue
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2)})
            continue

        if not stripped:
            flush_paragraph()
            continue

        # Skip metadata line
        if stripped.startswith("**Format**:"):
            continue

        # [Source: ...] lines attach to the preceding paragraph with a line break
        if stripped.startswith("[Source:"):
            if paragraph_lines:
                # Append source to current paragraph with a line break
                paragraph_lines.append("\\" )  # force line break
                paragraph_lines.append(stripped)
            elif blocks and blocks[-1]["type"] == "paragraph":
                # Previous paragraph already flushed — append source as line break
                blocks[-1]["runs"].append({"text": "\n", "bold": False, "italic": False, "url": None})
                blocks[-1]["runs"].extend(parse_inline(stripped))
            else:
                # No preceding paragraph — make it standalone
                blocks.append({"type": "paragraph", "runs": parse_inline(stripped)})
            continue

        # Numbered list items get a line break before them
        if re.match(r"^\d+\.\s", stripped) and paragraph_lines:
            paragraph_lines.append("\\")  # force line break

        paragraph_lines.append(stripped)

    flush_paragraph()
    return blocks


def add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from lxml import etree

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = etree.SubElement(paragraph._p, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), r_id)

    run_elem = etree.SubElement(hyperlink, qn("w:r"))
    rPr = etree.SubElement(run_elem, qn("w:rPr"))
    color = etree.SubElement(rPr, qn("w:color"))
    color.set(qn("w:val"), "0563C1")
    underline = etree.SubElement(rPr, qn("w:u"))
    underline.set(qn("w:val"), "single")

    text_elem = etree.SubElement(run_elem, qn("w:t"))
    text_elem.text = text
    text_elem.set(qn("xml:space"), "preserve")


def create_word_doc(blocks: list[dict], output_path: str) -> None:
    """Create a Word document from parsed markdown blocks."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    HEADING_SIZES = {1: Pt(23), 2: Pt(16), 3: Pt(13)}

    doc = Document()

    # Set default font and line spacing
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.space_before = Pt(0)

    # Configure heading styles
    for level in (1, 2, 3):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Arial"
        heading_style.font.size = HEADING_SIZES[level]
        heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        heading_style.paragraph_format.line_spacing = 1.15
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(4)

    for block in blocks:
        if block["type"] == "heading":
            level = block["level"]
            heading_text = block["text"]
            clean_text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', heading_text)
            doc.add_heading(clean_text, level=level)

        elif block["type"] == "horizontal_rule":
            from docx.oxml.ns import qn
            from lxml import etree
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")

        elif block["type"] == "paragraph":
            p = doc.add_paragraph()

            for run_data in block["runs"]:
                text = run_data["text"]
                parts = text.split("\n")
                for i, part in enumerate(parts):
                    if run_data["url"]:
                        add_hyperlink(p, part, run_data["url"])
                    elif part:
                        run = p.add_run(part)
                        if run_data["bold"]:
                            run.bold = True
                        if run_data["italic"]:
                            run.italic = True
                    if i < len(parts) - 1:
                        from docx.oxml.ns import qn as _qn
                        from lxml import etree as _etree
                        br_run = p.add_run()
                        _etree.SubElement(br_run._r, _qn("w:br"))

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Google Drive upload
# ---------------------------------------------------------------------------

def upload_to_google_drive(creds, docx_path: str, title: str) -> dict:
    """Upload .docx to Google Drive, converting to Google Doc."""
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from googleapiclient.errors import HttpError

    try:
        service = build("drive", "v3", credentials=creds)
    except HttpError as e:
        fail(f"Failed to connect to Google Drive API: {e}")

    file_metadata = {
        "name": title,
        "mimeType": "application/vnd.google-apps.document",
    }
    media = MediaFileUpload(
        docx_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    try:
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id",
        ).execute()
    except HttpError as e:
        fail(f"Failed to upload to Google Drive: {e}")

    doc_id = file.get("id")
    return {
        "success": True,
        "doc_url": f"https://docs.google.com/document/d/{doc_id}/edit",
        "doc_id": doc_id,
        "title": title,
    }


def extract_title_from_markdown(content: str) -> str:
    """Extract the first # heading from markdown, or return a default."""
    for line in content.split("\n"):
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return match.group(1)
    return "Untitled Script"


def main() -> None:
    if len(sys.argv) < 2 or sys.argv[1] in ("--help", "-h"):
        print_help()
        sys.exit(0)

    check_dependencies()

    script_path = sys.argv[1]
    title = None
    credentials_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "..", "..", "credentials.json"
    )

    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--title" and i + 1 < len(args):
            title = args[i + 1]
            i += 2
        elif args[i] == "--credentials" and i + 1 < len(args):
            credentials_path = args[i + 1]
            i += 2
        else:
            fail(f"Unknown argument: {args[i]}. Run with --help for usage.")

    if not os.path.exists(script_path):
        fail(f"Script file not found: {script_path}")

    with open(script_path, "r", encoding="utf-8") as f:
        content = f.read()

    if not content.strip():
        fail(f"Script file is empty: {script_path}")

    if title is None:
        title = extract_title_from_markdown(content)

    # Step 1: Parse markdown
    blocks = parse_markdown(content)

    # Step 2: Create temporary .docx with proper formatting
    tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_docx.close()
    try:
        create_word_doc(blocks, tmp_docx.name)

        # Step 3: Authenticate and upload to Google Drive (auto-converts to Google Doc)
        creds = authenticate(credentials_path)
        result = upload_to_google_drive(creds, tmp_docx.name, title)
    finally:
        os.unlink(tmp_docx.name)

    print(json.dumps(result))


if __name__ == "__main__":
    main()
